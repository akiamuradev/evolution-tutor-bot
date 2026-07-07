"""Small shared helper functions for bot handlers."""
import asyncio
import os
import time
from datetime import datetime

from aiogram.exceptions import TelegramBadRequest


UNLIMITED_USERS = [int(x.strip()) for x in os.getenv("UNLIMITED_USERS", "").split(",") if x.strip().isdigit()]
ADMIN_IDS = [int(x.strip()) for x in os.getenv("ADMIN_IDS", "").split(",") if x.strip().isdigit()]


def create_docx_file(text: str, title: str = "ЭВО:ЛЮЦИЯ") -> bytes:
    from docx import Document
    import io

    try:
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(text)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except Exception:
        return None


def create_pdf_file(text: str, title: str = "ЭВО:ЛЮЦИЯ") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    import io

    try:
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica", 12)
        y = height - 50
        for line in text.split("\n"):
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(50, y, line[:90])
            y -= 20
        c.save()
        buffer.seek(0)
        return buffer.read()
    except Exception:
        return None


def is_admin(tg_id: int) -> bool:
    return tg_id in ADMIN_IDS


def check_subscription(user: dict) -> bool:
    now = datetime.now()
    if user.get("tg_id") in UNLIMITED_USERS:
        return True
    if user.get("student_sub") and user.get("student_expires") and user["student_expires"].replace(tzinfo=None) > now:
        return True
    return False


def get_user_grade(user: dict) -> str:
    return user.get("grade_range") or "5-9"


async def show_thinking_animation(bot, chat_id, duration: int = 120):
    dots = [".", "..", "...", "...."]
    try:
        msg = await bot.send_message(chat_id, "🤔 Думаю")
        start_time = time.time()
        i = 0
        while time.time() - start_time < duration:
            await asyncio.sleep(0.8)
            i = (i + 1) % len(dots)
            try:
                await msg.edit_text(f"🤔 Думаю{dots[i]}")
            except TelegramBadRequest as e:
                if "message is not modified" in str(e):
                    continue
                break
            except Exception:
                break
        try:
            await msg.edit_text("⏳ Ещё немного...")
        except Exception:
            pass
        return msg
    except Exception:
        return None
